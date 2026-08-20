"""Document decoding and text extraction helpers."""

from __future__ import annotations

import re
from html import unescape
from urllib.parse import urljoin, urlparse

from .models import is_valid_calendar_date


def _decode_entities(value: str) -> str:
    return unescape(value).replace("\xa0", " ")


def decode_body(body: bytes, content_type: str) -> str:
    """Decode a response body using its declared charset when possible."""
    match = re.search(r"charset=[\"']?([\w-]+)", content_type, flags=re.I)
    charset = (match.group(1).lower() if match else "utf-8") or "utf-8"
    if charset == "gb2312":
        charset = "gbk"
    try:
        return body.decode(charset)
    except (LookupError, UnicodeDecodeError):
        return body.decode("utf-8", errors="replace")


def publication_date(raw: str | None) -> str | None:
    """Normalize a supported date string to ISO date form."""
    if not raw:
        return None
    match = re.search(r"(\d{4})[-/年.](\d{1,2})[-/月.](\d{1,2})", raw)
    if not match:
        return None
    value = (
        f"{match.group(1)}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"
    )
    return value if is_valid_calendar_date(value) else None


def extract_html(html: str) -> dict:
    """Extract normalized text and publication metadata from HTML."""
    title_match = re.search(
        r"<title[^>]*>([\s\S]*?)</title>", html, flags=re.I
    )
    title = " ".join(
        _decode_entities(title_match.group(1) if title_match else "").split()
    ).strip()
    meta_names = "article:published_time|pubdate|publishdate|dc\\.date|datepublished|article:modified_time|created|firstpublishedtime"
    meta = re.search(
        rf'<meta[^>]+(?:property|name)=["\'](?:{meta_names})["\'][^>]+content=["\']([^"\']+)',
        html,
        flags=re.I,
    )
    if not meta:
        meta = re.search(
            rf'<meta[^>]+content=["\']([^"\']+)["\'][^>]+(?:property|name)=["\'](?:{meta_names})["\']',
            html,
            flags=re.I,
        )
    time_match = re.search(
        r"<time[^>]+datetime=[\"']([^\"']+)", html, flags=re.I
    )
    meta_date = publication_date(meta.group(1) if meta else None)
    time_date = publication_date(time_match.group(1) if time_match else None)
    cleaned = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.I)
    cleaned = re.sub(r"<style[\s\S]*?</style>", " ", cleaned, flags=re.I)
    cleaned = re.sub(r"<noscript[\s\S]*?</noscript>", " ", cleaned, flags=re.I)
    cleaned = re.sub(r"<svg[\s\S]*?</svg>", " ", cleaned, flags=re.I)
    cleaned = re.sub(r"<(?:br|hr)\s*/?>", "\n", cleaned, flags=re.I)
    cleaned = re.sub(
        r"</(?:p|div|h[1-6]|li|tr|section|article|blockquote|table|ul|ol|header|footer|figure|time)>",
        "\n",
        cleaned,
        flags=re.I,
    )
    cleaned = re.sub(r"<[^>]+>", " ", cleaned)
    text = "\n".join(
        line.strip()
        for line in re.sub(
            r"[ \t]+", " ", _decode_entities(cleaned).replace("\r\n", "\n")
        ).split("\n")
        if line.strip()
    )
    if meta_date:
        publish_time, publish_time_source = meta_date, "meta"
    elif time_date:
        publish_time, publish_time_source = time_date, "time-element"
    else:
        text_date = None
        for match in re.finditer(
            r"(?:20\d{2}[-/年.]\d{1,2}[-/月.]\d{1,2})", text
        ):
            candidate = publication_date(match.group(0))
            if candidate:
                text_date = candidate
                break
        publish_time = text_date
        publish_time_source = "unknown"
    return {
        "title": title,
        "text": text,
        "publish_time": publish_time,
        "publish_time_source": publish_time_source,
    }


def extract_outbound_links(
    html: str, base_url: str, limit: int = 30
) -> list[dict]:
    """Extract unique external HTTP links from an HTML document."""
    seen: set[str] = set()
    links: list[dict] = []
    for match in re.finditer(
        r'<a[^>]+href=["\']([^"\']+)["\']', html, flags=re.I
    ):
        href = match.group(1).strip()
        if not href or href.startswith(
            ("#", "javascript:", "mailto:", "tel:")
        ):
            continue
        try:
            url = urljoin(base_url, href)
            parsed = urlparse(url)
        except Exception:
            continue
        if parsed.scheme not in ("http", "https") or parsed.hostname is None:
            continue
        if parsed.hostname == urlparse(base_url).hostname:
            continue
        normalized = url.split("#")[0].rstrip("/")
        if normalized in seen:
            continue
        seen.add(normalized)
        links.append({"url": normalized, "hostname": parsed.hostname.lower()})
        if len(links) >= limit:
            break
    return links


def extract_pdf_text(raw_bytes: bytes) -> str:
    """Extract normalized text from every page of a PDF document."""
    import fitz

    document = fitz.open(stream=raw_bytes, filetype="pdf")
    try:
        pages = [str(page.get_text("text")) for page in document]
    finally:
        document.close()
    text = "\n".join(pages)
    return "\n".join(line.strip() for line in text.split("\n") if line.strip())


def extract_docx_text(raw_bytes: bytes) -> str:
    """Extract paragraph and table text from a Word document."""
    from io import BytesIO

    from docx import Document

    document = Document(BytesIO(raw_bytes))
    parts = [
        str(part.text) for part in document.paragraphs if part.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip() for cell in row.cells if cell.text.strip()
            ]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
