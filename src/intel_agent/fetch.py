"""Secure document fetching with DNS pinning (port of fetch.ts)."""

from __future__ import annotations

import asyncio
import re
import ssl
from dataclasses import dataclass, field
from pathlib import Path
from typing import Awaitable, Callable
from urllib.parse import urlparse

from .models import IntelDocument, IntelError, is_valid_calendar_date
from .security import AddressResolver, resolve_public_url, source_group_of
from .source import source_type_for_domain
from .storage import read_json, sha256, verify_document_integrity, write_file_atomic, write_json_atomic

DEFAULT_MAX_BYTES = 5 * 1024 * 1024
DEFAULT_TIMEOUT_MS = 25_000
REDIRECT_STATUSES = {301, 302, 303, 307, 308}
MAX_REDIRECTS = 5


@dataclass
class FetchedResponse:
    status: int
    headers: dict[str, str] = field(default_factory=dict)
    body: bytes = b""


FetchLike = Callable[[str, dict | None, str], Awaitable[FetchedResponse]]


def _decode_entities(value: str) -> str:
    value = (
        value.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#39;", "'")
        .replace("&apos;", "'")
        .replace("&nbsp;", " ")
        .replace("&#160;", " ")
    )

    def dec(m: re.Match) -> str:
        try:
            return chr(int(m.group(1)))
        except ValueError:
            return ""

    value = re.sub(r"&#(\d+);", dec, value)

    def hex_dec(m: re.Match) -> str:
        try:
            return chr(int(m.group(1), 16))
        except ValueError:
            return ""

    return re.sub(r"&#x([0-9a-f]+);", hex_dec, value, flags=re.I)


def decode_body(body: bytes, content_type: str) -> str:
    m = re.search(r"charset=[\"']?([\w-]+)", content_type, flags=re.I)
    charset = (m.group(1).lower() if m else "utf-8") or "utf-8"
    if charset == "gb2312":
        charset = "gbk"
    try:
        return body.decode(charset)
    except (LookupError, UnicodeDecodeError):
        return body.decode("utf-8", errors="replace")


def publication_date(raw: str | None) -> str | None:
    if not raw:
        return None
    m = re.search(r"(\d{4})[-/年.](\d{1,2})[-/月.](\d{1,2})", raw)
    if not m:
        return None
    value = f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    return value if is_valid_calendar_date(value) else None


def extract_html(html: str) -> dict:
    title_m = re.search(r"<title[^>]*>([\s\S]*?)</title>", html, flags=re.I)
    title = " ".join(_decode_entities(title_m.group(1) if title_m else "").split()).strip()
    meta_names = "article:published_time|pubdate|publishdate|dc\\.date|datepublished|article:modified_time|created|firstpublishedtime"
    meta = re.search(
        rf'<meta[^>]+(?:property|name)=["\'](?:{meta_names})["\'][^>]+content=["\']([^"\']+)',
        html,
        flags=re.I,
    )
    if not meta:
        meta = re.search(
            rf'<meta[^>]+content=["\']([^"\']+)["\'][^>]+(?:property|name)=["\'](?:{meta_names})["\']',
            html,
            flags=re.I,
        )
    time_m = re.search(r"<time[^>]+datetime=[\"']([^\"']+)", html, flags=re.I)
    meta_date = publication_date(meta.group(1) if meta else None)
    time_date = publication_date(time_m.group(1) if time_m else None)
    cleaned = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.I)
    cleaned = re.sub(r"<style[\s\S]*?</style>", " ", cleaned, flags=re.I)
    cleaned = re.sub(r"<noscript[\s\S]*?</noscript>", " ", cleaned, flags=re.I)
    cleaned = re.sub(r"<svg[\s\S]*?</svg>", " ", cleaned, flags=re.I)
    cleaned = re.sub(r"<(?:br|hr)\s*/?>", "\n", cleaned, flags=re.I)
    cleaned = re.sub(
        r"</(?:p|div|h[1-6]|li|tr|section|article|blockquote|table|ul|ol|header|footer|figure|time)>",
        "\n",
        cleaned,
        flags=re.I,
    )
    cleaned = re.sub(r"<[^>]+>", " ", cleaned)
    text = "\n".join(
        line.strip()
        for line in re.sub(r"[ \t]+", " ", _decode_entities(cleaned).replace("\r\n", "\n")).split("\n")
        if line.strip()
    )
    if meta_date:
        publish_time, publish_time_source = meta_date, "meta"
    elif time_date:
        publish_time, publish_time_source = time_date, "time-element"
    else:
        text_date = None
        for m in re.finditer(r"(?:20\d{2}[-/年.]\d{1,2}[-/月.]\d{1,2})", text):
            candidate = publication_date(m.group(0))
            if candidate:
                text_date = candidate
                break
        publish_time = text_date
        publish_time_source = "unknown"
    return {"title": title, "text": text, "publish_time": publish_time, "publish_time_source": publish_time_source}


def extract_outbound_links(html: str, base_url: str, limit: int = 30) -> list[dict]:
    """提取 HTML 中的外链（供种子文档链接展开，绕过搜索直接扩展来源）。"""
    from urllib.parse import urljoin

    seen: set[str] = set()
    links: list[dict] = []
    for m in re.finditer(r'<a[^>]+href=["\']([^"\']+)["\']', html, flags=re.I):
        href = m.group(1).strip()
        if not href or href.startswith(("#", "javascript:", "mailto:", "tel:")):
            continue
        try:
            url = urljoin(base_url, href)
            parsed = urlparse(url)
        except Exception:
            continue
        if parsed.scheme not in ("http", "https") or parsed.hostname is None:
            continue
        if parsed.hostname == urlparse(base_url).hostname:
            continue
        normalized = url.split("#")[0].rstrip("/")
        if normalized in seen:
            continue
        seen.add(normalized)
        links.append({"url": normalized, "hostname": parsed.hostname.lower()})
        if len(links) >= limit:
            break
    return links


def extract_pdf_text(raw_bytes: bytes) -> str:
    """用 PyMuPDF 提取 PDF 全文（逐页合并，页间换行）。"""
    import fitz

    document = fitz.open(stream=raw_bytes, filetype="pdf")
    try:
        pages = [page.get_text("text") for page in document]
    finally:
        document.close()
    text = "\n".join(pages)
    return "\n".join(line.strip() for line in text.split("\n") if line.strip())


def extract_docx_text(raw_bytes: bytes) -> str:
    """用 python-docx 提取 .docx 全文（段落 + 表格）。"""
    from io import BytesIO

    from docx import Document

    document = Document(BytesIO(raw_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def canonicalize_url(raw: str) -> str:
    parsed = urlparse(raw)
    from urllib.parse import parse_qsl, urlencode

    params = [(k, v) for k, v in parse_qsl(parsed.query) if not re.match(r"^(?:utm_|spm$|from$|source$)", k, flags=re.I)]
    params.sort()
    query = urlencode(params)
    return parsed._replace(fragment="", query=query).geturl()


def injection_warnings(text: str) -> list[str]:
    patterns = [
        re.compile(r"忽略(?:此前|之前|以上).{0,12}指令", re.I),
        re.compile(r"(?:执行|运行).{0,8}(?:命令|代码)", re.I),
        re.compile(r"(?:调用|使用).{0,8}(?:工具|tool)", re.I),
        re.compile(r"ignore (?:all )?(?:previous|prior) instructions", re.I),
    ]
    return [line[:160] for line in text.split("\n") if any(p.search(line) for p in patterns)]


async def pinned_fetch(input_url: str, _init: dict | None, address: str, max_bytes: int = DEFAULT_MAX_BYTES) -> FetchedResponse:
    """DNS-pinned TCP/TLS fetch：连到已校验的公网 IP，TLS 仍用原主机名 SNI。"""
    parsed = urlparse(input_url)
    https = parsed.scheme == "https"
    port = parsed.port or (443 if https else 80)
    ssl_ctx = ssl.create_default_context() if https else None
    reader, writer = await asyncio.open_connection(
        address,
        port,
        ssl=ssl_ctx,
        server_hostname=parsed.hostname if https else None,
    )
    try:
        path = f"{parsed.path or '/'}{('?' + parsed.query) if parsed.query else ''}"
        request = (
            f"GET {path} HTTP/1.1\r\n"
            f"Host: {parsed.hostname}{f':{port}' if port not in (80, 443) else ''}\r\n"
            "User-Agent: pi-intelligence-collector/1.0\r\n"
            "Accept: text/html,text/plain,application/xhtml+xml\r\n"
            "Connection: close\r\n\r\n"
        )
        writer.write(request.encode("latin-1"))
        await writer.drain()
        chunks: list[bytes] = []
        size = 0
        while True:
            chunk = await reader.read(65_536)
            if not chunk:
                break
            size += len(chunk)
            if size > max_bytes + 65_536:
                raise IntelError("RESPONSE_TOO_LARGE", f"响应超过 {max_bytes} 字节")
            chunks.append(chunk)
    finally:
        writer.close()
        try:
            await writer.wait_closed()
        except Exception:
            pass
    return parse_http_response(b"".join(chunks), max_bytes)


async def httpx_fallback_fetch(input_url: str, _init: dict | None, _address: str, max_bytes: int = DEFAULT_MAX_BYTES) -> FetchedResponse:
    """httpx 回退抓取：兼容 WAF/Cloudflare 站点（无 DNS pinning，仅用于 IR/财报等低风险页）。"""
    import httpx

    async with httpx.AsyncClient(timeout=20.0, follow_redirects=False) as client:
        response = await client.get(
            input_url,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/pdf,text/plain,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.9,zh-CN;q=0.8",
            },
        )
        body = response.content
        if len(body) > max_bytes:
            raise IntelError("RESPONSE_TOO_LARGE", f"响应超过 {max_bytes} 字节")
        return FetchedResponse(
            status=response.status_code,
            headers={k.lower(): v for k, v in response.headers.items()},
            body=body,
        )


def parse_http_response(raw: bytes, max_bytes: int = DEFAULT_MAX_BYTES) -> FetchedResponse:
    separator = raw.find(b"\r\n\r\n")
    if separator < 0:
        raise IntelError("NETWORK_ERROR", "响应头不完整")
    lines = raw[:separator].decode("latin-1").split("\r\n")
    status_m = re.match(r"^HTTP/\d\.\d (\d{3})", lines[0])
    if not status_m:
        raise IntelError("NETWORK_ERROR", "响应状态无效")
    headers: dict[str, str] = {}
    for line in lines[1:]:
        if ":" in line:
            key, value = line.split(":", 1)
            headers[key.strip().lower()] = value.strip()
    body = raw[separator + 4 :]
    if re.search(r"chunked", headers.get("transfer-encoding", ""), re.I):
        body = decode_chunked(body)
    if len(body) > max_bytes:
        raise IntelError("RESPONSE_TOO_LARGE", f"响应超过 {max_bytes} 字节")
    return FetchedResponse(status=int(status_m.group(1)), headers=headers, body=body)


def decode_chunked(raw: bytes) -> bytes:
    chunks: list[bytes] = []
    offset = 0
    while offset < len(raw):
        line_end = raw.find(b"\r\n", offset)
        if line_end < 0:
            raise IntelError("NETWORK_ERROR", "分块响应不完整")
        size = int(raw[offset:line_end].split(b";", 1)[0].decode("ascii"), 16)
        offset = line_end + 2
        if size == 0:
            return b"".join(chunks)
        if offset + size + 2 > len(raw):
            raise IntelError("NETWORK_ERROR", "分块响应不完整")
        chunks.append(raw[offset : offset + size])
        offset += size + 2
    raise IntelError("NETWORK_ERROR", "分块响应缺少结束标记")


async def fetch_with_validated_redirects(
    raw_url: str,
    fetcher: FetchLike,
    resolver: AddressResolver | None,
    max_bytes: int,
) -> tuple[FetchedResponse, object]:
    current_url, addresses = await resolve_public_url(raw_url, resolver)
    for redirects in range(MAX_REDIRECTS + 1):
        response = await fetcher(_url_string(current_url), None, addresses[0])
        if response.status not in REDIRECT_STATUSES:
            return response, current_url
        if redirects >= MAX_REDIRECTS:
            raise IntelError("NETWORK_ERROR", "重定向次数超过限制")
        location = response.headers.get("location")
        if not location:
            raise IntelError("NETWORK_ERROR", "重定向响应缺少 Location")
        from urllib.parse import urljoin

        current_url, addresses = await resolve_public_url(urljoin(_url_string(current_url), location), resolver)
    raise IntelError("NETWORK_ERROR", "重定向次数超过限制")


def _url_string(parsed) -> str:
    return parsed.geturl() if hasattr(parsed, "geturl") else str(parsed)


async def fetch_document(
    cwd: Path,
    raw_url: str,
    fetcher: FetchLike | None = None,
    resolver: AddressResolver | None = None,
    max_bytes: int | None = None,
) -> tuple[IntelDocument, str, list[dict]]:
    max_bytes = max_bytes or DEFAULT_MAX_BYTES
    fetcher = fetcher or (lambda u, i, a: pinned_fetch(u, i, a, max_bytes))
    try:
        async with asyncio.timeout(DEFAULT_TIMEOUT_MS / 1000):
            response, final_url = await fetch_with_validated_redirects(raw_url, fetcher, resolver, max_bytes)
    except TimeoutError:
        raise IntelError("NETWORK_ERROR", f"抓取超时: {raw_url}")
    except IntelError:
        raise
    except Exception as error:
        raise IntelError("NETWORK_ERROR", f"抓取失败: {error}")
    if response.status != 200:
        raise IntelError("NETWORK_ERROR", f"抓取失败: HTTP {response.status}")
    content_type = response.headers.get("content-type", "").lower()
    allowed = r"^(?:text/html|application/xhtml\+xml|text/plain|application/pdf|application/msword|application/vnd\.openxmlformats-officedocument\.wordprocessingml\.document)(?:;|$)"
    if not re.match(allowed, content_type):
        raise IntelError("UNSUPPORTED_CONTENT", f"不支持的内容类型: {content_type or 'unknown'}")
    raw_bytes = response.body
    if len(raw_bytes) > max_bytes:
        raise IntelError("RESPONSE_TOO_LARGE", f"响应超过 {max_bytes} 字节")
    raw_text = decode_body(raw_bytes, content_type)
    is_html = bool(re.search(r"html|xhtml", content_type))
    is_pdf = "pdf" in content_type
    is_docx = "officedocument" in content_type or "msword" in content_type
    if is_html:
        extracted = extract_html(raw_text)
    elif is_pdf:
        try:
            text = extract_pdf_text(raw_bytes)
        except Exception as error:
            raise IntelError("NETWORK_ERROR", f"PDF 文本提取失败: {error}")
        extracted = {"title": Path(urlparse(_url_string(final_url)).path).name, "text": text, "publish_time": None, "publish_time_source": "unknown"}
    elif is_docx:
        try:
            text = extract_docx_text(raw_bytes)
        except Exception as error:
            raise IntelError("NETWORK_ERROR", f"Word 文本提取失败: {error}")
        extracted = {"title": Path(urlparse(_url_string(final_url)).path).name, "text": text, "publish_time": None, "publish_time_source": "unknown"}
    else:
        extracted = {"title": _url_string(final_url), "text": raw_text.replace("\r\n", "\n").strip(), "publish_time": None, "publish_time_source": "unknown"}
    outbound_links = extract_outbound_links(raw_text, _url_string(final_url)) if is_html else []
    if not extracted["publish_time"]:
        url_match = re.search(r"/(20\d{2})[-/]?(\d{2})[-/]?(\d{2})/", _url_string(final_url))
        if url_match and is_valid_calendar_date(f"{url_match.group(1)}-{url_match.group(2)}-{url_match.group(3)}"):
            extracted["publish_time"] = f"{url_match.group(1)}-{url_match.group(2)}-{url_match.group(3)}"
            extracted["publish_time_source"] = "unknown"
    canonical_url = canonicalize_url(_url_string(final_url))
    raw_hash = sha256(raw_bytes)
    document_id = f"doc-{sha256(f'{canonical_url}\n{raw_hash}')[:16]}"
    document_path = f"documents/{document_id}.json"
    if (cwd / "data/intel" / document_path).exists():
        existing = IntelDocument.model_validate(read_json(cwd, document_path))
        verify_document_integrity(cwd, existing)
        return existing, f"<untrusted_web_content>\n{extracted['text']}\n</untrusted_web_content>", outbound_links
    raw_path = f"data/raw/{document_id}.raw"
    text_path = f"data/raw/{document_id}.txt"
    write_file_atomic(cwd, raw_path, raw_bytes)
    write_file_atomic(cwd, text_path, extracted["text"])
    hostname = urlparse(_url_string(final_url)).hostname or ""
    document = IntelDocument(
        id=document_id,
        requested_url=raw_url,
        final_url=_url_string(final_url),
        canonical_url=canonical_url,
        title=extracted["title"] or _url_string(final_url),
        content_type=content_type.split(";")[0],
        publish_time=extracted["publish_time"],
        publish_time_source=extracted["publish_time_source"],
        collected_at=_now(),
        source_type=source_type_for_domain(hostname),
        source_group=source_group_of(_url_string(final_url)),
        raw_path=raw_path,
        raw_sha256=raw_hash,
        text_path=text_path,
        text_sha256=sha256(extracted["text"]),
        injection_warnings=injection_warnings(extracted["text"]),
    )
    write_json_atomic(cwd, document_path, document.model_dump())
    return document, f"<untrusted_web_content>\n{extracted['text']}\n</untrusted_web_content>", outbound_links


def _now() -> str:
    from datetime import datetime, timezone

    return datetime.now(timezone.utc).isoformat()
