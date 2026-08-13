"""Crawl resource extraction tests."""

from __future__ import annotations

import os
import signal
import subprocess
import sys
import threading
import time
import zipfile
from io import BytesIO
from pathlib import Path

import pytest

import intel_agent.extract as extract_module
from intel_agent.extract import extract_resource, is_rejected_resource


def _processor_error_whisper_worker(_audio, model_name, results):
    results.put((model_name, "processor error"))


def _large_result_whisper_worker(_audio, _model_name, results):
    results.put(
        (
            "complete",
            [(1.5, 62.25, "large transcript " * 100_000)],
        )
    )


def test_extract_html_text_links_and_plain_text():
    html = b'<html><body>Hello <a href="/next">next</a></body></html>'
    result = extract_resource(html, "text/html", "https://example.com/root")
    assert result.status == "complete"
    assert "Hello" in result.text
    assert result.links == ["https://example.com/next"]

    plain = extract_resource(
        b"name,value\na,1\n", "text/csv", "https://example.com/a.csv"
    )
    assert plain.text == "name,value\na,1"


def test_extract_html_discovers_embedded_media_and_metadata():
    html = b"""
    <html><head><title>Quarterly media briefing</title>
    <meta property="article:published_time" content="2026-08-12"></head>
    <body>
      <img src="/chart.webp" alt="orders chart">
      <audio controls src="/briefing.mp3"></audio>
      <video poster="/cover.jpg"><source src="/launch.webm"></video>
      <object data="/appendix.pdf"></object>
      <embed src="/slides.pptx">
    </body></html>
    """

    result = extract_resource(
        html,
        "text/html",
        "https://example.com/news/briefing",
        relevance_terms=["orders"],
    )

    assert result.title == "Quarterly media briefing"
    assert result.publish_time == "2026-08-12"
    assert result.publish_time_source == "meta"
    assert result.links == [
        "https://example.com/chart.webp",
        "https://example.com/briefing.mp3",
        "https://example.com/cover.jpg",
        "https://example.com/launch.webm",
        "https://example.com/appendix.pdf",
        "https://example.com/slides.pptx",
    ]
    assert result.link_relevance["https://example.com/chart.webp"] == 1


def test_pdf_relative_links_resolve_against_document_url(monkeypatch):
    class Page:
        def get_text(self, _kind):
            return "enough useful document text to avoid OCR"

        def get_links(self):
            return [{"uri": "../appendix.pdf"}]

    class Document:
        def __iter__(self):
            return iter([Page()])

        def __len__(self):
            return 1

        def close(self):
            pass

    class Fitz:
        @staticmethod
        def open(**_kwargs):
            return Document()

    monkeypatch.setitem(sys.modules, "fitz", Fitz)

    result = extract_resource(
        b"pdf",
        "application/pdf",
        "https://example.com/reports/annual.pdf",
    )

    assert result.status == "complete"
    assert result.links == ["https://example.com/appendix.pdf"]


def test_pdf_link_relevance_uses_document_and_link_context(monkeypatch):
    class Page:
        def get_text(self, _kind):
            return "quarterly outlook and enough useful document text"

        def get_links(self):
            return [
                {"uri": "https://example.com/appendix"},
                {"uri": "https://example.com/outlook"},
            ]

    class Document:
        def __iter__(self):
            return iter([Page()])

        def __len__(self):
            return 1

        def close(self):
            pass

    class Fitz:
        @staticmethod
        def open(**_kwargs):
            return Document()

    monkeypatch.setitem(sys.modules, "fitz", Fitz)

    result = extract_resource(
        b"pdf",
        "application/pdf",
        "https://example.com/report.pdf",
        relevance_terms=["outlook"],
    )

    assert result.link_relevance == {
        "https://example.com/appendix": 1,
        "https://example.com/outlook": 1,
    }


def test_duplicate_html_anchor_keeps_strongest_context():
    result = extract_resource(
        (
            b'<html><a href="/same">important topic progress</a>'
            b'<a href="/same">more</a></html>'
        ),
        "text/html",
        "https://example.com/root",
        relevance_terms=["important", "topic", "progress"],
    )

    assert result.link_relevance["https://example.com/same"] == 3


def test_html_navigation_links_are_ranked_below_article_links():
    result = extract_resource(
        (
            b"<html><nav><a href='/index'>topic report</a></nav>"
            b"<main><a href='/article'>topic report</a></main></html>"
        ),
        "text/html",
        "https://example.com/root",
        relevance_terms=["topic"],
    )

    assert (
        result.link_relevance["https://example.com/article"]
        > result.link_relevance["https://example.com/index"]
    )


def test_empty_html_is_unavailable_but_still_discovers_attachments():
    result = extract_resource(
        b'<html><body><a href="/report.pdf">report</a></body></html>',
        "text/html",
        "https://example.com/root",
    )

    assert result.status == "unavailable"
    assert result.links == ["https://example.com/report.pdf"]
    assert result.error == "html text is empty"


def test_plain_text_link_relevance_orders_links_from_document_context():
    result = extract_resource(
        (
            b"Appendix https://example.com/appendix\n"
            b"Market outlook https://example.com/outlook\n"
        ),
        "text/plain",
        "https://example.com/report.txt",
        relevance_terms=["outlook"],
    )

    assert (
        result.link_relevance["https://example.com/outlook"]
        > (result.link_relevance["https://example.com/appendix"])
    )


def test_ooxml_member_count_is_rejected_before_expansion():
    raw = BytesIO()
    with zipfile.ZipFile(raw, "w", zipfile.ZIP_DEFLATED) as archive:
        for number in range(1_025):
            archive.writestr(f"member-{number}.xml", b"x")

    result = extract_resource(
        raw.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "https://example.com/report.docx",
    )

    assert result.status == "failed"
    assert "OOXML member limit" in (result.error or "")


def test_ooxml_expanded_size_is_rejected_before_expansion():
    raw = BytesIO()
    with zipfile.ZipFile(raw, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("large.xml", b"0" * (32 * 1024 * 1024 + 1))

    result = extract_resource(
        raw.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "https://example.com/report.docx",
    )

    assert result.status == "failed"
    assert "OOXML expanded byte limit" in (result.error or "")


def test_pdf_page_count_is_rejected_before_page_work():
    import fitz

    document = fitz.open()
    for _ in range(201):
        document.new_page()
    raw = document.tobytes()
    document.close()

    result = extract_resource(
        raw, "application/pdf", "https://example.com/large.pdf"
    )

    assert result.status == "failed"
    assert "PDF page limit" in (result.error or "")


def test_pdf_raster_page_pixels_are_bounded_before_rendering():
    import fitz

    document = fitz.open()
    document.new_page(width=3_000, height=3_000)
    raw = document.tobytes()
    document.close()

    result = extract_resource(
        raw, "application/pdf", "https://example.com/huge-page.pdf"
    )

    assert result.status == "failed"
    assert "PDF raster pixel limit" in (result.error or "")


def test_image_pixels_are_bounded_before_ocr():
    import struct
    import zlib

    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + kind
            + data
            + struct.pack(">I", zlib.crc32(kind + data))
        )

    raw = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", 7_000, 7_000, 8, 2, 0, 0, 0))
        + chunk(b"IEND", b"")
    )

    result = extract_resource(raw, "image/png", "https://example.com/huge.png")

    assert result.status == "failed"
    assert "image pixel limit" in (result.error or "")


def test_heavy_extraction_timeout_terminates_its_worker(monkeypatch):
    process_state = {"alive": True, "terminated": False}

    class Results:
        def get(self, timeout):
            raise extract_module.queue.Empty

        def close(self):
            pass

        def cancel_join_thread(self):
            pass

    class Process:
        exitcode = None
        pid = None

        def __init__(self, **_kwargs):
            pass

        def start(self):
            pass

        def is_alive(self):
            return process_state["alive"]

        def join(self, timeout=None):
            pass

        def terminate(self):
            process_state.update(alive=False, terminated=True)

        def kill(self):
            process_state.update(alive=False, terminated=True)

    class Context:
        def Queue(self):
            return Results()

        def Process(self, **kwargs):
            return Process(**kwargs)

    monkeypatch.setattr(
        extract_module.multiprocessing, "get_context", lambda _: Context()
    )
    monkeypatch.setattr(extract_module, "_EXTRACTION_TIMEOUT_SECONDS", 0)
    result = extract_module.extract_resource_process(
        b"pdf", "application/pdf", "https://example.com/report.pdf"
    )

    assert result.status == "failed"
    assert result.error == "extraction timed out"
    assert process_state["terminated"] is True


def test_heavy_extraction_cancellation_terminates_parser_children(
    tmp_path, monkeypatch
):
    executable = tmp_path / "tesseract"
    pid_path = tmp_path / "parser.pid"
    executable.write_text(
        "#!/usr/bin/env python3\n"
        "import os\n"
        "from pathlib import Path\n"
        "import time\n"
        "Path(os.environ['INTEL_AGENT_TEST_PID_PATH']).write_text("
        "str(os.getpid()))\n"
        "time.sleep(30)\n",
        encoding="utf-8",
    )
    executable.chmod(0o755)
    monkeypatch.setenv("INTEL_AGENT_TEST_PID_PATH", str(pid_path))
    monkeypatch.setenv("PATH", f"{tmp_path}:{os.environ['PATH']}")
    cancellation = threading.Event()
    outcome: list[BaseException] = []

    def run() -> None:
        try:
            extract_module.extract_resource_process(
                b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
                b"\x00\x00\x00\x01\x00\x00\x00\x01",
                "image/png",
                "https://example.com/image.png",
                cancellation_event=cancellation,
            )
        except BaseException as error:
            outcome.append(error)

    worker = threading.Thread(target=run)
    worker.start()
    deadline = time.monotonic() + 5
    while not pid_path.exists() and time.monotonic() < deadline:
        time.sleep(0.01)
    assert pid_path.exists()
    pid = int(pid_path.read_text(encoding="utf-8"))

    def parser_is_alive() -> bool:
        try:
            os.kill(pid, 0)
        except ProcessLookupError:
            return False
        return True

    try:
        cancellation.set()
        worker.join(timeout=5)
        deadline = time.monotonic() + 2
        while parser_is_alive() and time.monotonic() < deadline:
            time.sleep(0.01)

        assert not worker.is_alive()
        assert outcome
        assert not parser_is_alive()
    finally:
        if parser_is_alive():
            os.kill(pid, signal.SIGKILL)


def test_ocr_text_has_line_numbers(monkeypatch):
    monkeypatch.setattr(
        "intel_agent.extract._ocr_image",
        lambda raw, languages: "first\nsecond",
    )
    result = extract_resource(
        b"image", "image/png", "https://example.com/image.png"
    )
    assert result.text == "1: first\n2: second"
    assert result.processor == "tesseract"


def test_transcript_has_timestamps(monkeypatch):
    monkeypatch.setattr(
        "intel_agent.extract._transcribe_media",
        lambda raw, suffix, model: [(0.0, 1.25, "hello")],
    )
    result = extract_resource(
        b"audio", "audio/mpeg", "https://example.com/clip.mp3"
    )
    assert result.text == "[00:00:00.000 --> 00:00:01.250] hello"
    assert result.processor == "faster-whisper"


def test_extract_xlsx_cells_and_hyperlinks_with_optional_processor_fake(
    monkeypatch,
):
    class Hyperlink:
        target = "https://example.com/external"

    class Cell:
        def __init__(self, coordinate, value, hyperlink=None):
            self.coordinate = coordinate
            self.value = value
            self.hyperlink = hyperlink

    class Sheet:
        title = "Summary"

        def iter_rows(self):
            return [
                [Cell("A1", "Report https://example.com/in-text")],
                [Cell("B2", 42, Hyperlink())],
            ]

    class Workbook:
        worksheets = [Sheet()]

        def close(self):
            pass

    class Openpyxl:
        @staticmethod
        def load_workbook(_stream, read_only, data_only):
            return Workbook()

    monkeypatch.setattr(extract_module, "import_module", lambda _: Openpyxl)

    result = extract_resource(
        b"workbook", "application/octet-stream", "https://example.com/a.xlsx"
    )

    assert result.status == "complete"
    assert result.processor == "openpyxl"
    assert result.text == (
        "Summary!A1: Report https://example.com/in-text\nSummary!B2: 42"
    )
    assert result.links == [
        "https://example.com/external",
        "https://example.com/in-text",
    ]


def test_xlsx_link_relevance_uses_visible_cell_context(monkeypatch):
    class Hyperlink:
        target = "https://example.com/outlook"

    class Cell:
        coordinate = "A1"
        value = "Market outlook"
        hyperlink = Hyperlink()

    class Sheet:
        title = "Summary"
        max_row = 1
        max_column = 1

        def iter_rows(self):
            return [[Cell()]]

    class Workbook:
        worksheets = [Sheet()]

        def close(self):
            pass

    class Openpyxl:
        @staticmethod
        def load_workbook(_stream, read_only, data_only):
            return Workbook()

    monkeypatch.setattr(extract_module, "import_module", lambda _: Openpyxl)

    result = extract_resource(
        b"workbook",
        "application/octet-stream",
        "https://example.com/a.xlsx",
        relevance_terms=["outlook"],
    )

    assert result.link_relevance["https://example.com/outlook"] > 0


def test_xlsx_sheet_count_is_bounded(monkeypatch):
    class Workbook:
        worksheets = [object()] * 101

        def close(self):
            pass

    class Openpyxl:
        @staticmethod
        def load_workbook(_stream, read_only, data_only):
            return Workbook()

    monkeypatch.setattr(extract_module, "import_module", lambda _: Openpyxl)

    result = extract_resource(
        b"workbook", "application/octet-stream", "https://example.com/a.xlsx"
    )

    assert result.status == "failed"
    assert "spreadsheet sheet limit" in (result.error or "")


def test_xlsx_declared_cell_work_is_bounded(monkeypatch):
    class Sheet:
        max_row = 501
        max_column = 501

        def iter_rows(self):
            raise AssertionError("oversized sheet must not be iterated")

    class Workbook:
        worksheets = [Sheet()]

        def close(self):
            pass

    class Openpyxl:
        @staticmethod
        def load_workbook(_stream, read_only, data_only):
            return Workbook()

    monkeypatch.setattr(extract_module, "import_module", lambda _: Openpyxl)

    result = extract_resource(
        b"workbook", "application/octet-stream", "https://example.com/a.xlsx"
    )

    assert result.status == "failed"
    assert "spreadsheet cell limit" in (result.error or "")


def test_extract_pptx_text_and_hyperlinks_with_optional_processor_fake(
    monkeypatch,
):
    class Hyperlink:
        address = "https://example.com/slide-link"

    class ClickAction:
        hyperlink = Hyperlink()

    class Shape:
        text = "Briefing https://example.com/in-text"
        click_action = ClickAction()

    class Slide:
        shapes = [Shape()]

    class Presentation:
        slides = [Slide()]

    class Pptx:
        @staticmethod
        def Presentation(_stream):
            return Presentation()

    monkeypatch.setattr(extract_module, "import_module", lambda _: Pptx)

    result = extract_resource(
        b"presentation",
        "application/octet-stream",
        "https://example.com/a.pptx",
    )

    assert result.status == "complete"
    assert result.processor == "python-pptx"
    assert result.text == "Slide 1: Briefing https://example.com/in-text"
    assert result.links == [
        "https://example.com/slide-link",
        "https://example.com/in-text",
    ]


def test_pptx_slide_count_is_bounded(monkeypatch):
    class Presentation:
        slides = [object()] * 501

    class Pptx:
        @staticmethod
        def Presentation(_stream):
            return Presentation()

    monkeypatch.setattr(extract_module, "import_module", lambda _: Pptx)

    result = extract_resource(
        b"presentation",
        "application/octet-stream",
        "https://example.com/a.pptx",
    )

    assert result.status == "failed"
    assert "presentation slide limit" in (result.error or "")


def test_legacy_doc_uses_generated_docx_content(monkeypatch):
    from docx import Document

    output = BytesIO()
    document = Document()
    document.add_paragraph(
        "Converted document https://example.com/document-link"
    )
    document.save(output)

    def run_process(command, _cancellation_event=None):
        directory = Path(command[command.index("--outdir") + 1])
        target = command[command.index("--convert-to") + 1]
        (directory / f"source.{target}").write_bytes(output.getvalue())
        return subprocess.CompletedProcess(command, 0, b"", b"")

    monkeypatch.setattr(
        extract_module.shutil, "which", lambda _: "libreoffice"
    )
    monkeypatch.setattr(extract_module, "_run_process", run_process)

    result = extract_resource(
        b"legacy document", "application/msword", "https://example.com/a.doc"
    )

    assert result.status == "complete"
    assert result.processor == "python-docx"
    assert "Converted document" in result.text
    assert result.links == ["https://example.com/document-link"]


def test_legacy_xls_uses_generated_xlsx_content(monkeypatch):
    class Hyperlink:
        target = "https://example.com/sheet-link"

    class Cell:
        coordinate = "A1"
        value = "Converted spreadsheet"
        hyperlink = Hyperlink()

    class Sheet:
        title = "Converted"

        def iter_rows(self):
            return [[Cell()]]

    class Workbook:
        worksheets = [Sheet()]

        def close(self):
            pass

    class EmptyWorkbook:
        worksheets: list[object] = []

        def close(self):
            pass

    class Openpyxl:
        @staticmethod
        def load_workbook(stream, read_only, data_only):
            return (
                Workbook()
                if stream.read() == b"converted xlsx"
                else EmptyWorkbook()
            )

    def run_process(command, _cancellation_event=None):
        directory = Path(command[command.index("--outdir") + 1])
        (directory / "source.xlsx").write_bytes(b"converted xlsx")
        return subprocess.CompletedProcess(command, 0, b"", b"")

    monkeypatch.setattr(
        extract_module.shutil, "which", lambda _: "libreoffice"
    )
    monkeypatch.setattr(extract_module, "_run_process", run_process)
    monkeypatch.setattr(extract_module, "import_module", lambda _: Openpyxl)

    result = extract_resource(
        b"legacy workbook",
        "application/vnd.ms-excel",
        "https://example.com/a.xls",
    )

    assert result.status == "complete"
    assert result.processor == "openpyxl"
    assert result.text == "Converted!A1: Converted spreadsheet"
    assert result.links == ["https://example.com/sheet-link"]


def test_legacy_ppt_uses_generated_pptx_content(monkeypatch):
    class Hyperlink:
        address = "https://example.com/presentation-link"

    class ClickAction:
        hyperlink = Hyperlink()

    class Shape:
        text = "Converted presentation"
        click_action = ClickAction()

    class Slide:
        shapes = [Shape()]

    class Presentation:
        slides = [Slide()]

    class EmptyPresentation:
        slides: list[object] = []

    class Pptx:
        @staticmethod
        def Presentation(stream):
            return (
                Presentation()
                if stream.read() == b"converted pptx"
                else EmptyPresentation()
            )

    def run_process(command, _cancellation_event=None):
        directory = Path(command[command.index("--outdir") + 1])
        (directory / "source.pptx").write_bytes(b"converted pptx")
        return subprocess.CompletedProcess(command, 0, b"", b"")

    monkeypatch.setattr(
        extract_module.shutil, "which", lambda _: "libreoffice"
    )
    monkeypatch.setattr(extract_module, "_run_process", run_process)
    monkeypatch.setattr(extract_module, "import_module", lambda _: Pptx)

    result = extract_resource(
        b"legacy presentation",
        "application/vnd.ms-powerpoint",
        "https://example.com/a.ppt",
    )

    assert result.status == "complete"
    assert result.processor == "python-pptx"
    assert result.text == "Slide 1: Converted presentation"
    assert result.links == ["https://example.com/presentation-link"]


def test_whisper_worker_returns_timestamped_segments_from_processor(
    monkeypatch,
):
    class Segment:
        start = 1.25
        end = 2.5
        text = " spoken text "

    class WhisperModel:
        def __init__(self, model_name):
            self.model_name = model_name

        def transcribe(self, _audio_path):
            return [Segment()], object()

    faster_whisper = type("FasterWhisper", (), {"WhisperModel": WhisperModel})

    class Results:
        values: list[tuple[str, object]] = []

        def put(self, value):
            self.values.append(value)

    monkeypatch.setattr(
        extract_module, "import_module", lambda _: faster_whisper
    )
    results = Results()

    extract_module._whisper_worker("audio.wav", "small", results)

    assert results.values == [("complete", [(1.25, 2.5, "spoken text")])]


@pytest.mark.parametrize(
    ("error", "expected_status"),
    [
        (ImportError("faster-whisper missing"), "unavailable"),
        (RuntimeError("model failed"), "failed"),
    ],
)
def test_whisper_worker_returns_processor_errors(
    monkeypatch, error, expected_status
):
    def unavailable(_module_name):
        raise error

    class Results:
        values: list[tuple[str, object]] = []

        def put(self, value):
            self.values.append(value)

    monkeypatch.setattr(extract_module, "import_module", unavailable)
    results = Results()

    extract_module._whisper_worker("audio.wav", "small", results)

    assert results.values == [(expected_status, str(error))]


def test_whisper_runner_returns_complete_worker_result(monkeypatch, tmp_path):
    class Results:
        def __init__(self):
            self.values = [("complete", [(0, 1.5, "transcript")])]

        def get(self, timeout):
            return self.values.pop(0)

        def close(self):
            pass

        def cancel_join_thread(self):
            pass

    class Process:
        exitcode = 0

        def __init__(self, **_kwargs):
            pass

        def start(self):
            pass

        def is_alive(self):
            return False

        def join(self, timeout=None):
            pass

        def terminate(self):
            pass

        def kill(self):
            pass

    class Context:
        def Queue(self):
            return Results()

        def Process(self, **kwargs):
            return Process(**kwargs)

    monkeypatch.setattr(
        extract_module.multiprocessing, "get_context", lambda _: Context()
    )

    segments = extract_module._run_whisper_worker(
        tmp_path / "audio.wav", "small", None
    )

    assert segments == [(0.0, 1.5, "transcript")]


@pytest.mark.parametrize(
    ("worker_result", "expected_status"),
    [("unavailable", "unavailable"), ("failed", "failed")],
)
def test_whisper_worker_preserves_processor_failure_semantics(
    monkeypatch, worker_result, expected_status
):
    def fake_run_process(command, _cancellation_event=None):
        Path(command[-1]).write_bytes(b"wav")
        return subprocess.CompletedProcess(command, 0, b"", b"")

    monkeypatch.setattr(
        extract_module, "_whisper_worker", _processor_error_whisper_worker
    )
    monkeypatch.setattr(extract_module, "_run_process", fake_run_process)
    monkeypatch.setattr(extract_module.shutil, "which", lambda _name: "ffmpeg")

    result = extract_resource(
        b"audio",
        "audio/mpeg",
        "https://example.com/clip.mp3",
        whisper_model=worker_result,
    )

    assert result.status == expected_status
    assert result.error == "processor error"


def test_large_whisper_result_completes_without_queue_deadlock(monkeypatch):
    def fake_run_process(command, _cancellation_event=None):
        Path(command[-1]).write_bytes(b"wav")
        return subprocess.CompletedProcess(command, 0, b"", b"")

    monkeypatch.setattr(
        extract_module, "_whisper_worker", _large_result_whisper_worker
    )
    monkeypatch.setattr(extract_module, "_run_process", fake_run_process)
    monkeypatch.setattr(extract_module.shutil, "which", lambda _name: "ffmpeg")
    cancellation = threading.Event()
    finished = threading.Event()
    outcome = []

    def run() -> None:
        try:
            outcome.append(
                extract_resource(
                    b"audio",
                    "audio/mpeg",
                    "https://example.com/clip.mp3",
                    cancellation_event=cancellation,
                )
            )
        finally:
            finished.set()

    extraction = threading.Thread(target=run)
    extraction.start()
    completed_without_cancellation = finished.wait(timeout=8)
    if not completed_without_cancellation:
        cancellation.set()
        assert finished.wait(timeout=2), "Whisper worker cleanup also blocked"
    extraction.join()

    assert completed_without_cancellation, "large Whisper result deadlocked"
    assert outcome[0].status == "complete"
    expected_segment = "large transcript " * 100_000
    assert outcome[0].text == (
        f"[00:00:01.500 --> 00:01:02.250] {expected_segment}"
    )


def test_missing_tesseract_executable_marks_extraction_unavailable(
    monkeypatch,
):
    def unavailable(raw, languages):
        raise FileNotFoundError("tesseract")

    monkeypatch.setattr("intel_agent.extract._ocr_image", unavailable)
    result = extract_resource(
        b"original", "image/jpeg", "https://example.com/photo.jpg"
    )
    assert result.status == "unavailable"
    assert "tesseract" in (result.error or "")


def test_malformed_image_oserror_marks_extraction_failed(monkeypatch):
    def malformed(raw, languages):
        raise OSError("cannot identify image file")

    monkeypatch.setattr("intel_agent.extract._ocr_image", malformed)
    result = extract_resource(
        b"malformed", "image/png", "https://example.com/photo.png"
    )
    assert result.status == "failed"
    assert "cannot identify" in (result.error or "")


@pytest.mark.parametrize(
    ("mime_type", "url"),
    [
        ("image/svg+xml", "https://example.com/image.svg"),
        ("image/gif", "https://example.com/image.gif"),
        ("audio/aac", "https://example.com/audio.aac"),
        ("video/x-msvideo", "https://example.com/video.avi"),
    ],
)
def test_rejects_unlisted_media_formats(mime_type, url):
    result = extract_resource(b"original", mime_type, url)
    assert result.status == "skipped"


@pytest.mark.parametrize(
    ("mime_type", "url"),
    [
        ("application/zip", "https://example.com/a.zip"),
        ("application/javascript", "https://example.com/a.js"),
        ("application/x-msdownload", "https://example.com/a.exe"),
        ("text/plain", "https://example.com/image.svg"),
    ],
)
def test_rejects_archives_scripts_and_executables(mime_type, url):
    assert is_rejected_resource(mime_type, url)


def test_owned_subprocess_terminates_when_extraction_is_cancelled(tmp_path):
    run_process = getattr(extract_module, "_run_process", None)
    assert callable(run_process)
    pid_path = tmp_path / "child.pid"
    cancelled = threading.Event()
    outcome: list[BaseException] = []

    def run() -> None:
        try:
            run_process(
                [
                    sys.executable,
                    "-c",
                    (
                        "import os,time,pathlib;"
                        f"pathlib.Path({str(pid_path)!r}).write_text(str(os.getpid()));"
                        "time.sleep(30)"
                    ),
                ],
                cancelled,
            )
        except BaseException as error:
            outcome.append(error)

    worker = threading.Thread(target=run)
    worker.start()
    deadline = time.monotonic() + 2
    while not pid_path.exists() and time.monotonic() < deadline:
        time.sleep(0.01)
    assert pid_path.exists()
    pid = int(pid_path.read_text())

    cancelled.set()
    worker.join(timeout=2)

    assert not worker.is_alive()
    assert outcome
    with pytest.raises(ProcessLookupError):
        os.kill(pid, 0)
