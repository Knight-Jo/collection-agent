"""PDF/Word document extraction and outbound link expansion tests."""

import pytest

from intel_agent.fetch import FetchedResponse, extract_docx_text, extract_outbound_links, extract_pdf_text, fetch_document

PDF_RAW = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length 60>>stream
BT /F1 18 Tf 72 720 Td (Low-altitude economy PDF report) Tj ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000052 00000 n 
0000000101 00000 n 
0000000223 00000 n 
0000000346 00000 n 
trailer<</Size 6/Root 1 0 R>>
startxref
405
%%EOF
"""


def test_extract_outbound_links():
    html = '<a href="https://www.gov.cn/a">gov</a><a href="https://news.cn/b">news</a><a href="#anchor">same</a><a href="/relative">rel</a><a href="https://www.gov.cn/a">dup</a>'
    links = extract_outbound_links(html, "https://example.com/page")
    urls = [l["url"] for l in links]
    assert "https://www.gov.cn/a" in urls
    assert "https://news.cn/b" in urls
    # 去重 + 过滤锚点 + 同域（含相对路径）过滤
    assert len(urls) == 2
    assert all(l["hostname"] for l in links)


def test_extract_outbound_links_filters_same_host():
    html = '<a href="https://example.com/other">self</a><a href="https://gov.cn/x">other</a>'
    links = extract_outbound_links(html, "https://example.com/page")
    assert len(links) == 1
    assert links[0]["hostname"] == "gov.cn"


def test_extract_pdf_text():
    text = extract_pdf_text(PDF_RAW)
    assert "Low-altitude economy PDF report" in text


def test_extract_docx_text():
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("亿航智能订单报告第一段")
    doc.add_paragraph("第二段数据 100 架")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "订单数"
    table.rows[0].cells[1].text = "1000"
    buf = BytesIO()
    doc.save(buf)
    text = extract_docx_text(buf.getvalue())
    assert "亿航智能订单报告第一段" in text
    assert "订单数" in text


@pytest.mark.asyncio
async def test_fetch_pdf_document(cwd):
    async def fetcher(url, init, address):
        return FetchedResponse(status=200, headers={"content-type": "application/pdf"}, body=PDF_RAW)

    async def resolver(hostname):
        return ["93.184.216.34"]

    document, content, links = await fetch_document(
        cwd, "https://report.example.com/low-altitude.pdf", fetcher=fetcher, resolver=resolver
    )
    assert "PDF report" in content
    assert document.content_type == "application/pdf"
    assert links == []


@pytest.mark.asyncio
async def test_fetch_html_returns_outbound_links(cwd):
    html = '<html><body><a href="https://gov.cn/next">next</a>正文</body></html>'

    async def fetcher(url, init, address):
        return FetchedResponse(status=200, headers={"content-type": "text/html"}, body=html.encode())

    async def resolver(hostname):
        return ["93.184.216.34"]

    document, content, links = await fetch_document(
        cwd, "https://news.example.com/a", fetcher=fetcher, resolver=resolver
    )
    assert len(links) == 1
    assert links[0]["url"] == "https://gov.cn/next"
